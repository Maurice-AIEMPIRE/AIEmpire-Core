"""
Legal Document Processing Pipeline
====================================
Handles the complete document intake pipeline for legal analysis:

1. Scan local directories for legal documents
2. Extract text from PDFs (native + OCR fallback)
3. Extract entities (persons, dates, amounts, companies)
4. Normalize file naming per convention
5. Deduplicate documents
6. Build structured inventory + tags

Designed to run on Hetzner server where documents are stored locally.
All processing happens locally (P3 privacy level).

Usage:
    from data.legal_doc_processor import LegalDocProcessor

    processor = LegalDocProcessor()
    result = await processor.process_inbox()

    # Or process a single file:
    doc = await processor.process_file("/path/to/document.pdf")

CLI:
    python -m data.legal_doc_processor scan         # Scan inbox
    python -m data.legal_doc_processor process      # Full processing
    python -m data.legal_doc_processor inventory    # Show inventory
    python -m data.legal_doc_processor stats        # Processing stats
"""

import asyncio
import hashlib
import json
import os
import re
import shutil
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

# ── Project paths ─────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).parent.parent
DATA_INBOX = PROJECT_ROOT / "data" / "inbox"
DATA_PROCESSED = PROJECT_ROOT / "data" / "processed"
DATA_INDEX = PROJECT_ROOT / "data" / "index"
INVENTORY_FILE = DATA_PROCESSED / "INVENTORY.json"
TAGS_FILE = DATA_INDEX / "TAGS.json"

# ── Supported extensions ──────────────────────────────────────────────
PDF_EXTENSIONS = {".pdf"}
OFFICE_EXTENSIONS = {".docx", ".doc", ".xlsx", ".xls", ".odt", ".rtf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp"}
TEXT_EXTENSIONS = {".txt", ".eml", ".msg", ".csv"}
ALL_LEGAL_EXTENSIONS = PDF_EXTENSIONS | OFFICE_EXTENSIONS | IMAGE_EXTENSIONS | TEXT_EXTENSIONS

# ── Document types for naming convention ──────────────────────────────
DOC_TYPE_KEYWORDS = {
    "CONTRACT": ["vertrag", "contract", "vereinbarung", "agreement", "nda"],
    "EMAIL": ["email", "e-mail", "mail", "nachricht", "message", "korrespondenz"],
    "INVOICE": ["rechnung", "invoice", "quittung", "receipt", "zahlung"],
    "LETTER": ["brief", "letter", "schreiben", "anschreiben", "mitteilung"],
    "TESTIMONY": ["aussage", "testimony", "zeugnis", "erklaerung", "declaration"],
    "REPORT": ["bericht", "report", "gutachten", "expertise", "stellungnahme"],
    "SCREENSHOT": ["screenshot", "bildschirm", "screen"],
    "CERTIFICATE": ["zertifikat", "certificate", "urkunde", "bescheinigung", "zeugnis"],
    "COURT": ["gericht", "court", "klage", "urteil", "beschluss", "verfuegung", "ladung"],
    "PROTOCOL": ["protokoll", "protocol", "niederschrift", "minutes"],
    "NOTICE": ["kuendigung", "abmahnung", "mahnung", "notice", "warning"],
    "PHOTO": ["foto", "photo", "bild", "image"],
}


@dataclass
class ExtractedEntity:
    """An entity extracted from a document."""
    entity_type: str  # PERSON, DATE, AMOUNT, COMPANY, REFERENCE
    value: str
    context: str = ""  # surrounding text
    confidence: str = "medium"  # low, medium, high


@dataclass
class ProcessedDocument:
    """Full metadata for a processed legal document."""
    original_path: str
    processed_path: str
    filename: str
    normalized_name: str
    file_type: str  # CONTRACT, EMAIL, INVOICE, etc.
    extension: str
    size_bytes: int
    sha256: str
    text_content: str = ""
    text_length: int = 0
    extraction_method: str = ""  # native, ocr, office, plain
    extraction_confidence: str = "medium"
    entities: list[ExtractedEntity] = field(default_factory=list)
    dates_found: list[str] = field(default_factory=list)
    persons_found: list[str] = field(default_factory=list)
    amounts_found: list[str] = field(default_factory=list)
    companies_found: list[str] = field(default_factory=list)
    language: str = "de"
    processed_at: str = ""
    page_count: int = 0
    is_duplicate: bool = False
    duplicate_of: str = ""


class LegalDocProcessor:
    """
    Processes legal documents through the full pipeline:
    scan -> extract -> entities -> normalize -> dedupe -> index.
    """

    def __init__(self, inbox_dir: Optional[Path] = None, processed_dir: Optional[Path] = None):
        self.inbox_dir = inbox_dir or DATA_INBOX
        self.processed_dir = processed_dir or DATA_PROCESSED
        self.index_dir = DATA_INDEX
        self._inventory: dict[str, dict] = {}
        self._tags: dict[str, dict] = {}
        self._load_state()

    def _load_state(self):
        """Load existing inventory and tags."""
        if INVENTORY_FILE.exists():
            try:
                with open(INVENTORY_FILE, "r") as f:
                    self._inventory = json.load(f)
            except (json.JSONDecodeError, OSError):
                self._inventory = {}

        if TAGS_FILE.exists():
            try:
                with open(TAGS_FILE, "r") as f:
                    self._tags = json.load(f)
            except (json.JSONDecodeError, OSError):
                self._tags = {}

    def _save_inventory(self):
        """Save inventory to disk (atomic write)."""
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        tmp = INVENTORY_FILE.with_suffix(".tmp")
        with open(tmp, "w") as f:
            json.dump(self._inventory, f, indent=2, ensure_ascii=False)
        tmp.replace(INVENTORY_FILE)

    def _save_tags(self):
        """Save tags to disk (atomic write)."""
        self.index_dir.mkdir(parents=True, exist_ok=True)
        tmp = TAGS_FILE.with_suffix(".tmp")
        with open(tmp, "w") as f:
            json.dump(self._tags, f, indent=2, ensure_ascii=False)
        tmp.replace(TAGS_FILE)

    def _compute_sha256(self, filepath: Path) -> str:
        """Compute SHA256 hash of a file."""
        h = hashlib.sha256()
        with open(filepath, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        return h.hexdigest()

    def scan_inbox(self) -> list[Path]:
        """Scan inbox directory for legal documents."""
        if not self.inbox_dir.exists():
            print(f"[WARN] Inbox-Verzeichnis existiert nicht: {self.inbox_dir}")
            return []

        files = []
        for path in sorted(self.inbox_dir.rglob("*")):
            if path.is_file() and path.suffix.lower() in ALL_LEGAL_EXTENSIONS:
                files.append(path)

        return files

    def scan_directory(self, directory: str | Path) -> list[Path]:
        """Scan any directory for legal documents (for NAS/Hetzner local paths)."""
        directory = Path(directory)
        if not directory.exists():
            print(f"[WARN] Verzeichnis existiert nicht: {directory}")
            return []

        files = []
        for path in sorted(directory.rglob("*")):
            if path.is_file() and path.suffix.lower() in ALL_LEGAL_EXTENSIONS:
                files.append(path)

        return files

    # ── Text Extraction ───────────────────────────────────────────────

    def _extract_pdf_text(self, filepath: Path) -> tuple[str, str, int]:
        """
        Extract text from PDF. Returns (text, method, page_count).
        Tries native extraction first, falls back to OCR.
        """
        text = ""
        method = "none"
        page_count = 0

        # Try pymupdf (fitz) first - best native PDF extraction
        try:
            import fitz  # pymupdf

            doc = fitz.open(str(filepath))
            page_count = len(doc)
            pages_text = []
            for page in doc:
                page_text = page.get_text()
                if page_text.strip():
                    pages_text.append(page_text)
            doc.close()

            if pages_text:
                text = "\n\n--- Seite ---\n\n".join(pages_text)
                method = "pymupdf"
        except ImportError:
            pass
        except Exception as e:
            print(f"  [WARN] pymupdf Fehler bei {filepath.name}: {e}")

        # If no text extracted, try pdfplumber
        if not text.strip():
            try:
                import pdfplumber

                with pdfplumber.open(str(filepath)) as pdf:
                    page_count = len(pdf.pages)
                    pages_text = []
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            pages_text.append(page_text)
                    if pages_text:
                        text = "\n\n--- Seite ---\n\n".join(pages_text)
                        method = "pdfplumber"
            except ImportError:
                pass
            except Exception as e:
                print(f"  [WARN] pdfplumber Fehler bei {filepath.name}: {e}")

        # If still no text, try OCR (for scanned documents)
        if not text.strip():
            text, method, page_count = self._extract_pdf_ocr(filepath, page_count)

        return text, method, page_count

    def _extract_pdf_ocr(self, filepath: Path, page_count: int) -> tuple[str, str, int]:
        """OCR fallback for scanned PDFs."""
        try:
            import fitz  # pymupdf for image extraction

            doc = fitz.open(str(filepath))
            if not page_count:
                page_count = len(doc)

            pages_text = []
            for page_num, page in enumerate(doc):
                # Render page to image
                pix = page.get_pixmap(dpi=300)
                img_bytes = pix.tobytes("png")

                # OCR with tesseract
                try:
                    from PIL import Image
                    import pytesseract
                    import io

                    img = Image.open(io.BytesIO(img_bytes))
                    page_text = pytesseract.image_to_string(img, lang="deu+eng")
                    if page_text.strip():
                        pages_text.append(page_text)
                except ImportError:
                    print("  [WARN] pytesseract/PIL nicht installiert. OCR uebersprungen.")
                    break

            doc.close()

            if pages_text:
                return "\n\n--- Seite (OCR) ---\n\n".join(pages_text), "ocr", page_count

        except ImportError:
            pass
        except Exception as e:
            print(f"  [WARN] OCR Fehler: {e}")

        return "", "failed", page_count

    def _extract_image_text(self, filepath: Path) -> tuple[str, str]:
        """Extract text from images via OCR."""
        try:
            from PIL import Image
            import pytesseract

            img = Image.open(str(filepath))
            text = pytesseract.image_to_string(img, lang="deu+eng")
            return text, "ocr"
        except ImportError:
            print("  [WARN] pytesseract/PIL nicht installiert.")
            return "", "failed"
        except Exception as e:
            print(f"  [WARN] Bild-OCR Fehler: {e}")
            return "", "failed"

    def _extract_office_text(self, filepath: Path) -> tuple[str, str]:
        """Extract text from Office documents (DOCX, XLSX, etc.)."""
        ext = filepath.suffix.lower()

        if ext == ".docx":
            try:
                from docx import Document

                doc = Document(str(filepath))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return "\n".join(paragraphs), "python-docx"
            except ImportError:
                print("  [WARN] python-docx nicht installiert.")
                return "", "failed"

        if ext in (".xlsx", ".xls"):
            try:
                import openpyxl

                wb = openpyxl.load_workbook(str(filepath), read_only=True)
                rows = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        row_text = " | ".join(str(c) for c in row if c is not None)
                        if row_text.strip():
                            rows.append(row_text)
                wb.close()
                return "\n".join(rows), "openpyxl"
            except ImportError:
                print("  [WARN] openpyxl nicht installiert.")
                return "", "failed"

        if ext == ".rtf":
            try:
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                # Basic RTF stripping
                text = re.sub(r'\\[a-z]+\d*\s?', '', content)
                text = re.sub(r'[{}]', '', text)
                return text.strip(), "rtf-strip"
            except Exception:
                return "", "failed"

        return "", "unsupported"

    def _extract_plain_text(self, filepath: Path) -> tuple[str, str]:
        """Extract text from plain text files."""
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                with open(filepath, "r", encoding=encoding) as f:
                    return f.read(), f"plain-{encoding}"
            except (UnicodeDecodeError, OSError):
                continue
        return "", "failed"

    def extract_text(self, filepath: Path) -> tuple[str, str, int]:
        """
        Extract text from any supported document.
        Returns (text, extraction_method, page_count).
        """
        ext = filepath.suffix.lower()
        page_count = 0

        if ext in PDF_EXTENSIONS:
            return self._extract_pdf_text(filepath)
        elif ext in IMAGE_EXTENSIONS:
            text, method = self._extract_image_text(filepath)
            return text, method, 1
        elif ext in OFFICE_EXTENSIONS:
            text, method = self._extract_office_text(filepath)
            return text, method, 0
        elif ext in TEXT_EXTENSIONS:
            text, method = self._extract_plain_text(filepath)
            return text, method, 0
        else:
            return "", "unsupported", 0

    # ── Entity Extraction ─────────────────────────────────────────────

    def extract_entities(self, text: str) -> list[ExtractedEntity]:
        """Extract legal-relevant entities from text."""
        entities = []

        # German dates: DD.MM.YYYY, DD.MM.YY, DD. Monat YYYY
        date_patterns = [
            r'\b(\d{1,2}\.\d{1,2}\.\d{4})\b',
            r'\b(\d{1,2}\.\d{1,2}\.\d{2})\b',
            r'\b(\d{1,2}\.\s*(?:Januar|Februar|Maerz|März|April|Mai|Juni|Juli|'
            r'August|September|Oktober|November|Dezember)\s*\d{4})\b',
            # ISO dates
            r'\b(\d{4}-\d{2}-\d{2})\b',
        ]
        for pattern in date_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                entities.append(ExtractedEntity(
                    entity_type="DATE",
                    value=match.group(1),
                    context=text[max(0, match.start()-30):match.end()+30],
                    confidence="high",
                ))

        # Euro amounts: EUR 1.234,56 or 1.234,56 EUR or 1.234,56 Euro
        amount_patterns = [
            r'(?:EUR|Euro|€)\s*(\d{1,3}(?:\.\d{3})*(?:,\d{2})?)',
            r'(\d{1,3}(?:\.\d{3})*(?:,\d{2})?)\s*(?:EUR|Euro|€)',
        ]
        for pattern in amount_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                entities.append(ExtractedEntity(
                    entity_type="AMOUNT",
                    value=match.group(0).strip(),
                    context=text[max(0, match.start()-30):match.end()+30],
                    confidence="high",
                ))

        # Legal references: § 123 BGB, Art. 5 GG, etc.
        ref_patterns = [
            r'(§§?\s*\d+[a-z]?\s*(?:Abs\.\s*\d+)?\s*(?:S\.\s*\d+)?\s*'
            r'(?:Nr\.\s*\d+)?\s*(?:BGB|HGB|StGB|ZPO|GG|ArbGG|BetrVG|'
            r'KSchG|AGG|BDSG|DSGVO|GewO|TzBfG|MuSchG|SGB|BUrlG|'
            r'ArbZG|ArbSchG|InsO|UWG|GmbHG|AktG|BauGB|VOB|HOAI|'
            r'DIN\s*\d+))',
            r'(Art\.\s*\d+\s*(?:Abs\.\s*\d+)?\s*(?:GG|EMRK|AEUV|EUV|GRCh))',
        ]
        for pattern in ref_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                entities.append(ExtractedEntity(
                    entity_type="LEGAL_REFERENCE",
                    value=match.group(1).strip(),
                    context=text[max(0, match.start()-30):match.end()+30],
                    confidence="high",
                ))

        # Company patterns: GmbH, AG, KG, etc.
        company_patterns = [
            r'([A-ZÄÖÜ][A-Za-zäöüß&\-\s]{1,40}\s+'
            r'(?:GmbH|AG|KG|OHG|GbR|e\.K\.|UG|SE|Co\.\s*KG|mbH|'
            r'Partnerschaft|eG|e\.V\.|Ltd\.|Inc\.|Corp\.))',
        ]
        for pattern in company_patterns:
            for match in re.finditer(pattern, text):
                company = match.group(1).strip()
                if len(company) > 3:  # Filter noise
                    entities.append(ExtractedEntity(
                        entity_type="COMPANY",
                        value=company,
                        context=text[max(0, match.start()-30):match.end()+30],
                        confidence="medium",
                    ))

        # Aktenzeichen (case numbers): Az. 1 BvR 123/20, 3 O 456/21
        case_patterns = [
            r'(?:Az\.?|Aktenzeichen)\s*:?\s*(\d+\s*[A-Za-z]+\s*\d+/\d{2,4})',
        ]
        for pattern in case_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                entities.append(ExtractedEntity(
                    entity_type="CASE_NUMBER",
                    value=match.group(1).strip(),
                    context=text[max(0, match.start()-30):match.end()+30],
                    confidence="high",
                ))

        return entities

    # ── Document Classification ───────────────────────────────────────

    def classify_document(self, filepath: Path, text: str) -> str:
        """Classify document type based on filename and content."""
        name_lower = filepath.stem.lower()
        text_lower = text[:2000].lower()  # Check first 2000 chars

        for doc_type, keywords in DOC_TYPE_KEYWORDS.items():
            for kw in keywords:
                if kw in name_lower or kw in text_lower:
                    return doc_type

        # Fallback based on extension
        if filepath.suffix.lower() in IMAGE_EXTENSIONS:
            return "PHOTO"

        return "OTHER"

    def normalize_filename(self, filepath: Path, doc_type: str, text: str) -> str:
        """
        Create normalized filename: YYYY-MM-DD_TYPE_DESCRIPTION.ext

        Tries to extract a date from the document content or filename.
        """
        # Try to find a date
        date_str = "UNDATED"

        # From filename
        date_match = re.search(r'(\d{4}[-_]\d{2}[-_]\d{2})', filepath.stem)
        if date_match:
            date_str = date_match.group(1).replace("_", "-")
        else:
            date_match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', filepath.stem)
            if date_match:
                date_str = f"{date_match.group(3)}-{date_match.group(2)}-{date_match.group(1)}"

        # From content (first date found)
        if date_str == "UNDATED" and text:
            content_date = re.search(r'(\d{1,2})\.(\d{1,2})\.(\d{4})', text[:3000])
            if content_date:
                day = content_date.group(1).zfill(2)
                month = content_date.group(2).zfill(2)
                year = content_date.group(3)
                date_str = f"{year}-{month}-{day}"

        # Clean description from filename
        desc = filepath.stem.lower()
        desc = re.sub(r'\d{4}[-_]\d{2}[-_]\d{2}', '', desc)  # Remove dates
        desc = re.sub(r'[^a-z0-9äöüß]+', '-', desc)  # Clean special chars
        desc = desc.strip("-")[:50]  # Limit length

        if not desc:
            desc = "dokument"

        ext = filepath.suffix.lower()
        return f"{date_str}_{doc_type}_{desc}{ext}"

    # ── Deduplication ─────────────────────────────────────────────────

    def check_duplicate(self, sha256: str) -> Optional[str]:
        """Check if a document with this hash already exists."""
        for path, meta in self._inventory.items():
            if meta.get("sha256") == sha256:
                return path
        return None

    # ── Main Processing ───────────────────────────────────────────────

    async def process_file(self, filepath: Path | str) -> ProcessedDocument:
        """Process a single legal document through the full pipeline."""
        filepath = Path(filepath)
        if not filepath.exists():
            raise FileNotFoundError(f"Datei nicht gefunden: {filepath}")

        print(f"  [PROCESS] {filepath.name} ...", end=" ", flush=True)

        # Basic metadata
        size = filepath.stat().st_size
        sha256 = self._compute_sha256(filepath)

        # Check for duplicates
        duplicate_of = self.check_duplicate(sha256)

        # Extract text
        text, method, page_count = self.extract_text(filepath)

        # Determine confidence
        if method in ("pymupdf", "pdfplumber", "python-docx", "openpyxl"):
            confidence = "high"
        elif method in ("ocr", "rtf-strip"):
            confidence = "medium"
        elif method == "failed":
            confidence = "low"
        else:
            confidence = "medium"

        # Classify document type
        doc_type = self.classify_document(filepath, text)

        # Normalize filename
        normalized_name = self.normalize_filename(filepath, doc_type, text)

        # Extract entities
        entities = self.extract_entities(text) if text else []

        # Categorize entities
        dates = [e.value for e in entities if e.entity_type == "DATE"]
        persons = [e.value for e in entities if e.entity_type == "PERSON"]
        amounts = [e.value for e in entities if e.entity_type == "AMOUNT"]
        companies = [e.value for e in entities if e.entity_type == "COMPANY"]

        # Copy to processed directory
        processed_path = self.processed_dir / normalized_name
        self.processed_dir.mkdir(parents=True, exist_ok=True)

        # Handle name collisions
        if processed_path.exists() and not duplicate_of:
            counter = 1
            stem = processed_path.stem
            ext = processed_path.suffix
            while processed_path.exists():
                processed_path = self.processed_dir / f"{stem}_{counter}{ext}"
                counter += 1

        if not duplicate_of:
            shutil.copy2(str(filepath), str(processed_path))

        doc = ProcessedDocument(
            original_path=str(filepath),
            processed_path=str(processed_path),
            filename=filepath.name,
            normalized_name=normalized_name,
            file_type=doc_type,
            extension=filepath.suffix.lower(),
            size_bytes=size,
            sha256=sha256,
            text_content=text,
            text_length=len(text),
            extraction_method=method,
            extraction_confidence=confidence,
            entities=entities,
            dates_found=dates,
            persons_found=persons,
            amounts_found=amounts,
            companies_found=companies,
            processed_at=datetime.now(timezone.utc).isoformat(),
            page_count=page_count,
            is_duplicate=bool(duplicate_of),
            duplicate_of=duplicate_of or "",
        )

        # Update inventory
        self._inventory[str(processed_path)] = {
            "original": str(filepath),
            "normalized_name": normalized_name,
            "type": doc_type,
            "size_bytes": size,
            "sha256": sha256,
            "text_length": len(text),
            "extraction_method": method,
            "confidence": confidence,
            "page_count": page_count,
            "dates": dates[:10],
            "amounts": amounts[:10],
            "companies": companies[:10],
            "is_duplicate": bool(duplicate_of),
            "processed_at": doc.processed_at,
        }

        # Update tags
        self._tags[str(processed_path)] = {
            "type": doc_type,
            "entities": [asdict(e) for e in entities[:50]],
            "dates": dates[:20],
            "amounts": amounts[:20],
            "companies": companies[:20],
            "legal_refs": [e.value for e in entities if e.entity_type == "LEGAL_REFERENCE"][:20],
            "case_numbers": [e.value for e in entities if e.entity_type == "CASE_NUMBER"][:10],
        }

        status = "DUP" if duplicate_of else "OK"
        print(f"{status} ({method}, {len(text)} Zeichen, {len(entities)} Entitaeten)")
        return doc

    async def process_inbox(self) -> list[ProcessedDocument]:
        """Process all files in the inbox directory."""
        files = self.scan_inbox()
        if not files:
            print("[INFO] Keine Dateien in data/inbox/ gefunden.")
            print(f"  Lege Dokumente in {self.inbox_dir} ab und starte erneut.")
            return []

        print(f"\n[PIPELINE] Verarbeite {len(files)} Dateien aus {self.inbox_dir}")
        print("=" * 60)

        results = []
        for filepath in files:
            try:
                doc = await self.process_file(filepath)
                results.append(doc)
            except Exception as e:
                print(f"  [ERROR] {filepath.name}: {e}")

        # Save state
        self._save_inventory()
        self._save_tags()

        # Print summary
        print("\n" + "=" * 60)
        print(f"[ERGEBNIS] {len(results)} von {len(files)} verarbeitet")
        types = {}
        for doc in results:
            types[doc.file_type] = types.get(doc.file_type, 0) + 1
        for doc_type, count in sorted(types.items()):
            print(f"  {doc_type}: {count}")

        dupes = sum(1 for d in results if d.is_duplicate)
        if dupes:
            print(f"  Duplikate: {dupes}")

        failed = sum(1 for d in results if d.extraction_method == "failed")
        if failed:
            print(f"  Extraktion fehlgeschlagen: {failed}")

        return results

    async def process_directory(self, directory: str | Path) -> list[ProcessedDocument]:
        """Process all legal documents in a given directory (NAS/Hetzner local path)."""
        files = self.scan_directory(directory)
        if not files:
            print(f"[INFO] Keine Dokumente in {directory} gefunden.")
            return []

        print(f"\n[PIPELINE] Verarbeite {len(files)} Dateien aus {directory}")
        print("=" * 60)

        results = []
        for filepath in files:
            try:
                doc = await self.process_file(filepath)
                results.append(doc)
            except Exception as e:
                print(f"  [ERROR] {filepath.name}: {e}")

        self._save_inventory()
        self._save_tags()
        return results

    def get_inventory(self) -> dict:
        """Get the current document inventory."""
        return self._inventory

    def get_stats(self) -> dict:
        """Get processing statistics."""
        total = len(self._inventory)
        types = {}
        total_text = 0
        total_entities = 0
        duplicates = 0

        for meta in self._inventory.values():
            doc_type = meta.get("type", "OTHER")
            types[doc_type] = types.get(doc_type, 0) + 1
            total_text += meta.get("text_length", 0)
            if meta.get("is_duplicate"):
                duplicates += 1

        for tag_data in self._tags.values():
            total_entities += len(tag_data.get("entities", []))

        return {
            "total_documents": total,
            "by_type": types,
            "total_text_chars": total_text,
            "total_entities": total_entities,
            "duplicates": duplicates,
            "unique_documents": total - duplicates,
        }

    def get_text_for_analysis(self, max_docs: int = 0) -> list[dict]:
        """
        Get all extracted texts ready for legal analysis.
        Returns list of {path, type, text, entities, dates, amounts, companies}.
        """
        docs = []
        for path, meta in self._inventory.items():
            if meta.get("is_duplicate"):
                continue

            # Load text from processed file or cached content
            text = ""
            processed_path = Path(path)
            if processed_path.suffix == ".txt" and processed_path.exists():
                try:
                    with open(processed_path, "r") as f:
                        text = f.read()
                except Exception:
                    pass

            tags = self._tags.get(path, {})
            docs.append({
                "path": path,
                "type": meta.get("type", "OTHER"),
                "text": text,
                "text_length": meta.get("text_length", 0),
                "confidence": meta.get("confidence", "unknown"),
                "dates": tags.get("dates", []),
                "amounts": tags.get("amounts", []),
                "companies": tags.get("companies", []),
                "legal_refs": tags.get("legal_refs", []),
                "case_numbers": tags.get("case_numbers", []),
            })

            if max_docs and len(docs) >= max_docs:
                break

        return docs


# ── CLI ───────────────────────────────────────────────────────────────
async def main():
    import sys

    processor = LegalDocProcessor()
    cmd = sys.argv[1] if len(sys.argv) > 1 else "process"

    if cmd == "scan":
        files = processor.scan_inbox()
        print(f"\n{len(files)} Dateien in {DATA_INBOX} gefunden:")
        for f in files:
            size_kb = f.stat().st_size / 1024
            print(f"  {f.name} ({size_kb:.1f} KB)")

    elif cmd == "process":
        if len(sys.argv) > 2:
            # Process specific directory
            await processor.process_directory(sys.argv[2])
        else:
            await processor.process_inbox()

    elif cmd == "inventory":
        inv = processor.get_inventory()
        if not inv:
            print("[INFO] Inventar leer. Erst 'process' ausfuehren.")
            return
        print(f"\n{len(inv)} Dokumente im Inventar:")
        for path, meta in inv.items():
            print(f"  [{meta.get('type', '?'):12}] {meta.get('normalized_name', path)}")
            print(f"    Konfidenz: {meta.get('confidence')} | "
                  f"Text: {meta.get('text_length', 0)} Zeichen | "
                  f"Seiten: {meta.get('page_count', 0)}")

    elif cmd == "stats":
        stats = processor.get_stats()
        print("\n=== Verarbeitungsstatistik ===")
        print(f"Dokumente gesamt: {stats['total_documents']}")
        print(f"Eindeutige: {stats['unique_documents']}")
        print(f"Duplikate: {stats['duplicates']}")
        print(f"Text extrahiert: {stats['total_text_chars']:,} Zeichen")
        print(f"Entitaeten gefunden: {stats['total_entities']}")
        print("\nNach Typ:")
        for doc_type, count in sorted(stats["by_type"].items()):
            print(f"  {doc_type}: {count}")

    else:
        print("Usage: python -m data.legal_doc_processor [scan|process|inventory|stats]")
        print("       python -m data.legal_doc_processor process /pfad/zu/dokumenten")


if __name__ == "__main__":
    asyncio.run(main())
